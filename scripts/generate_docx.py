from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()
    
    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    
    # Path to the MD file
    md_path = r'c:\Users\ketan\Workspace\argus-transaction-monitor\Project_Report_Argus.md'
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Main Title (h1)
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # Subtitle (h2)
            p = doc.add_heading(line[3:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('### '):
            # Section Heading (h3)
            doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            # Horizontal line (skip or add page break)
            doc.add_page_break()
        elif line.startswith('- '):
            # Bullet points
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**'):
            # Bold lines
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
        elif line.startswith('[Insert Screenshot]'):
            # Placeholder for screenshots
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = (150, 150, 150)
        else:
            # Normal paragraph
            doc.add_paragraph(line)
            
    # Save the document
    output_path = r'c:\Users\ketan\Workspace\argus-transaction-monitor\Project_Report_Argus.docx'
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

try:
    create_report()
except Exception as e:
    print(f"Error: {e}")
